import os
from pathlib import Path
from typing import Union, Dict, Any
import json
import re
from docx import Document as DocxDocument
import PyPDF2


class DocumentLoader:
    SUPPORTED_FORMATS = {
        '.md': 'markdown',
        '.txt': 'plaintext',
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.json': 'json',
        '.rst': 'rst',
        '.tex': 'latex'
    }

    @staticmethod
    def load(file_path: str) -> Dict[str, Any]:
        path = Path(file_path)
        extension = path.suffix.lower()
        format_type = DocumentLoader.SUPPORTED_FORMATS[extension]
        loader_method = getattr(DocumentLoader, f'_load_{format_type}')
        content = loader_method(file_path)
        
        return {
            'content': content,
            'file_path': file_path,
            'file_name': path.name,
            'format': format_type,
            'extension': extension,
            'encoding': 'utf-8'
        }

    @staticmethod
    def _load_markdown(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    @staticmethod
    def _load_plaintext(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        text_content = []
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            num_pages = len(pdf_reader.pages)
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                text_content.append(page.extract_text())
        
        return '\n'.join(text_content)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = DocxDocument(file_path)
        paragraphs = [paragraph.text for paragraph in doc.paragraphs]
        
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text for cell in row.cells]
                paragraphs.append('|'.join(row_cells))
        
        return '\n'.join(paragraphs)

    @staticmethod
    def _load_json(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return json.dumps(data, indent=2, ensure_ascii=False)

    @staticmethod
    def _load_rst(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    @staticmethod
    def _load_latex(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    @staticmethod
    def get_supported_formats() -> list:
        return list(DocumentLoader.SUPPORTED_FORMATS.keys())

    @staticmethod
    def auto_detect_language(content: str) -> str:
        if content.count('def ') > content.count('λ'):
            return 'en'
        elif content.count('def ') + content.count('class ') > 10:
            return 'en'
        elif '№' in content or 'ё' in content.lower():
            return 'ru'
        elif '你好' in content:
            return 'zh'
        elif 'التعريف' in content:
            return 'ar'
        else:
            return 'en'

    @staticmethod
    def extract_metadata(document_data: Dict[str, Any]) -> Dict[str, Any]:
        content = document_data['content']
        lines = content.split('\n')
        title = ''
        
        for line in lines[:20]:
            if line.startswith('# '):
                title = line.replace('# ', '').strip()
                break
            elif line.startswith('# '):
                title = line.replace('#', '').strip()
                break
        
        word_count = len(content.split())
        char_count = len(content)
        headings = [line for line in lines if line.startswith('#')]
        definition_count = len([l for l in lines if 'Definition' in l or 'Определение' in l])
        theorem_count = len([l for l in lines if 'Theorem' in l or 'Теорема' in l])
        
        return {
            'title': title or document_data['file_name'],
            'word_count': word_count,
            'char_count': char_count,
            'heading_count': len(headings),
            'definition_count': definition_count,
            'theorem_count': theorem_count,
            'language': DocumentLoader.auto_detect_language(content),
            'source_file': document_data['file_path']
        }
